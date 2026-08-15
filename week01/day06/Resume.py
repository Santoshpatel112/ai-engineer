# 1. Job Description Schema
# 2. Resume Schama
# 3. Pdf pdne ka Schema
# 4.Actial Read krna file ko. read Resume 
# 5.final Score 

import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field

load_dotenv()
my_api_key=os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("API key kaha hai bhai")

client=Groq(api_key=my_api_key)
model = "llama-3.1-8b-instant"


job_description="""
Job Description

Join our core engineering team to design, build, and optimize web applications that power our B2B SaaS platforms. You will contribute to real production systems, working across both frontend and backend stacks.

Key Responsibilities
Develop and enhance web application features using modern frontend and backend technologies
Integrate RESTful APIs and third-party services into existing and new product modules
Collaborate with the product and design teams to translate requirements into working features
Write clean, modular, and well-documented code following best practices
Debug and resolve application issues, optimizing performance where required
Participate in code reviews and contribute to technical discussions
Required Skills & Qualifications
Solid understanding of JavaScript, HTML5, and CSS3
Familiarity with React.js or similar frontend frameworks
Basic to intermediate knowledge of Node.js or equivalent backend technologies
Understanding of RESTful APIs, JSON, and HTTP protocols
Exposure to databases such as MongoDB or MySQL
Logical problem-solving and debugging ability
Good to Have
Prior internship or project experience in web development
Familiarity with version control tools such as Git
Basic understanding of cloud services or deployment concepts


"""
class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: str | None  # Changed from float to str to handle "15+ years"
    education_requirements: list[str]
    responsibilities: list[str]

jobd_schema = JobD.model_json_schema()

system_prompt = f"""
You are an expert HR assistant.

Your job is to analyze job descriptions and extract structured information from them.

Return ONLY valid JSON matching this schema:
{jobd_schema}

CRITICAL RULES:
1. Return ONLY the JSON object - no extra text
2. Do NOT return the schema itself
3. Use exact field names: role, required_skills, preferred_skills, minimum_experience, education_requirements, responsibilities
4. All arrays must contain only strings - no Python syntax like + or concatenation
5. For minimum_experience, use a string like "15+ years" or null if not specified
6. If information is missing, use empty array [] or null
7. Do not invent information

Example valid response:
{{
  "role": "Software Engineer",
  "required_skills": ["Python", "Java"],
  "preferred_skills": ["AWS", "Docker"],
  "minimum_experience": "5+ years",
  "education_requirements": ["Bachelor's degree in Computer Science"],
  "responsibilities": ["Develop software", "Debug issues"]
}}
"""

user_prompt = f"""
Analyze the following job description:

{job_description}
"""
message_system={
    "role" : "system",
    "content" : system_prompt
}
message_user={
    "role" : "user",
    "content" : user_prompt
}
response_format={
    "type" : "json_object"
}


messages=[message_system, message_user]

response=client.chat.completions.create(
    model=model, 
    messages=messages, 
    response_format=response_format,
    temperature=0
)


answer=response.choices[0].message.content

raw_json=answer
# print("Raw JSON response:")
# print(raw_json)
# print("\n" + "="*50 + "\n")



import json
job_data=json.loads(raw_json)

job = JobD(**job_data)

print(job.minimum_experience)
print(job.education_requirements)



#parse real
class MatchResult(BaseModel):
    score: float
    details: dict
class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []

class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None

    total_experience_years: float | None = None

    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []


resume_schema = Resume.model_json_schema()
def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    
    Return ONLY valid JSON matching this schema:
    {match_schema}

    Your response must be valid JSON with these exact fields:
    - score: number from 0 to 100
    - details: object containing analysis

    Include in details:
    - candidate_name: string
    - matching_skills: array of strings
    - missing_skills: array of strings  
    - experience_met: string (Yes/No with brief explanation)
    - overall_percentage: number from 0 to 100
    - summary: string with final verdict

    Ensure all strings are properly quoted and JSON is valid.
    """
    
    system_message = {
        "role": "system",
        "content": "You are an expert HR assistant. Return only valid JSON responses."
    }
    
    user_message = {
        "role": "user",
        "content": prompt
    }
    
    messages = [system_message, user_message]
    response_format = {
        "type": "json_object"
    }
    
    response = client.chat.completions.create(
        model=model, 
        messages=messages, 
        response_format=response_format,
        temperature=0  # For consistent JSON generation
    )
    
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)
def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume and return ONLY valid JSON matching this schema:

    {resume_schema}

    IMPORTANT:
    - Return a JSON OBJECT, not an array
    - Use exact field names from schema (lowercase with underscores)
    - If a value is not available, use null
    - If a list has no information, use empty array []
    - Include internships in experiences array
    - Extract skills from entire resume
    
    Example valid response:
    {{
        "name": "John Doe",
        "email": "john@email.com", 
        "phone": "123-456-7890",
        "total_experience_years": 5.0,
        "skills": ["Python", "Java"],
        "experiences": [{{
            "company": "ABC Corp",
            "role": "Developer", 
            "duration": "2020-2023",
            "description": "Built web apps",
            "skills_used": ["Python", "React"]
        }}],
        "education": ["Bachelor's in CS"],
        "projects": ["Web App Project"],
        "certifications": ["AWS Certified"]
    }}
    """
    
    user_prompt = f"""
    Parse the following resume text and return ONLY the JSON object:

    {resume_text}
    """
    
    message_system = {
        "role": "system",
        "content": system_prompt
    }
    message_user = {
        "role": "user",
        "content": user_prompt
    }
    messages = [message_system, message_user]
    response_format = {
        "type": "json_object"
    }
    
    response = client.chat.completions.create(
        model=model, 
        messages=messages, 
        response_format=response_format,
        temperature=0
    )
    
    raw_output = response.choices[0].message.content
    
    try:
        data = json.loads(raw_output)
        if isinstance(data, list):
            print("Warning: LLM returned list instead of object, using first item")
            data = data[0] if data else {}
        
        # Fix null values that should be empty lists
        for field in ['skills', 'experiences', 'education', 'projects', 'certifications']:
            if data.get(field) is None:
                data[field] = []
                
        resume = Resume(**data)
        return resume
    except Exception as e:
        print(f"Error parsing resume: {e}")
        print(f"Raw output: {raw_output}")
        # Return empty resume as fallback
        return Resume()


from pypdf import PdfReader
from docx import Document
def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None



# lets do it now - process just the first resume to avoid rate limits
resume_folder = Path("resumes")
all_results = []
processed_count = 0
max_resumes = 1  # Process only 1 resume to avoid rate limits

for file_path in resume_folder.iterdir():
    if processed_count >= max_resumes:
        break
        
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue
    
    print(f"\n🔄 Processing: {file_path.name}")
    
    try:
        resume_text = read_resume(file_path)
        if not resume_text or len(resume_text.strip()) < 50:
            print(f"❌ Could not read or text too short: {file_path.name}")
            continue
            
        print("📄 Resume text extracted successfully")
        
        # Truncate resume text if too long to avoid token limits
        if len(resume_text) > 3000:
            resume_text = resume_text[:3000] + "..."
            print("⚠️  Resume text truncated to avoid token limits")
        
        parsed_resume = parse_resume(resume_text)  # llm call1
        print("✅ Resume parsed successfully")
        
        time.sleep(3)
        
        result = final_score(job, parsed_resume)  # llm call2
        print(f"🎯 Score: {result.score}/100")
        
        all_results.append({
            "name": parsed_resume.name or file_path.stem,
            "score": result.score,
            "details": result.details
        })
        
        processed_count += 1
        time.sleep(3)
        
    except Exception as e:
        print(f"❌ Error processing {file_path.name}: {e}")
        continue
print("\n" + "="*60)
print("🏆 RESUME EVALUATION RESULTS")
print("="*60)

if all_results:
    all_results.sort(key=lambda candidate: candidate["score"], reverse=True)
    
    for i, candidate in enumerate(all_results, 1):
        print(f"\n#{i} {candidate['name']}")
        print(f"📊 Score: {candidate['score']}/100")
        print(f"📋 Details: {candidate['details']}")
        
else:
    print("No resumes were successfully processed.")

print("\n" + "="*60)
print("✨ Evaluation Complete!")
print("="*60)